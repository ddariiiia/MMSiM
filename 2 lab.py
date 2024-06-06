import bitstring
import docx

def to_binary(file, encoding_type):
    with open(file, "rb") as f:
        text = f.read().decode("UTF-8")
    result = ""
    for symb in text:
        result += bitstring.BitArray(symb.encode(encoding_type)).bin
    return result

def hide_message(binary_message, empty_name, filled_name):
    filled_container = docx.Document()
    empty_container = docx.Document(empty_name)
    container_text = "\n".join([p.text for p in empty_container.paragraphs])
    paragraph = filled_container.add_paragraph()
    for i in range(len(container_text)):
        if container_text[i] == "\n":
            paragraph = filled_container.add_paragraph()
            continue
        if i < len(binary_message) and binary_message[i] == "1":
            run = paragraph.add_run(container_text[i])
            run.font.size = docx.shared.Pt(13.5)
            run.font.name = "Helvetica"
            run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
        else:
            if run.font.size == docx.shared.Pt(22):  # Пропускаем символы, если размер шрифта равен 22
                continue
            run = paragraph.add_run(container_text[i])
            run.font.size = docx.shared.Pt(13)
            run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
            run.font.name = "Helvetica"
    filled_container.save(filled_name)

def decode_binary(binary_message, encoding_type):
    if encoding_type == "bo2":  # Декодируем как двоичные данные
        return bitstring.BitArray(bin=binary_message).bytes.decode('utf-8')
    else:
        hex_str = hex(int(binary_message, 2))[2:]
        decoded_message = bytes.fromhex(hex_str).decode(encoding_type)
        return decoded_message

if __name__ == "__main__":
    empty_container = "C://Users//daria//OneDrive//Рабочий стол//10 семестр//Мат медоды сокрытия и маскирования информации//22//empty.docx"
    filled_container = "C://Users//daria//OneDrive//Рабочий стол//10 семестр//Мат медоды сокрытия и маскирования информации//22//filled.docx"
    encoding_type = "windows-1251"
    message = "C://Users//daria//OneDrive//Рабочий стол//10 семестр//Мат медоды сокрытия и маскирования информации//22//message.txt"
    binary_message = to_binary(message, encoding_type)
    print("Скрываемое сообщение в двоичном виде: " + binary_message)
    hide_message(binary_message, empty_container, filled_container)
    decoded_message = decode_binary(binary_message, encoding_type)

    
    method = 'font_size'
    print(f"Использованный метод: {method}")

    encodings = ["windows-1251", "koi8-r", "cp866", "bo2"]
    for encoding in encodings:
        if encoding == "bo2":
            continue  
        decoded_message = decode_binary(binary_message, encoding)
        print(f"Декодированное сообщение для {encoding}: {decoded_message}")
