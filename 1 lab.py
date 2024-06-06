from docx import Document

def analyze_document(docx_file):
    document = Document(docx_file)
    methods = create_text_array_and_methods(document)
    chosen_method = determine_hiding_method(methods)
    encoded_text = encode_text(document, chosen_method)
    decoded_results = decode_text(encoded_text)
    print_results(decoded_results, chosen_method)


def create_text_array_and_methods(document):
    methods = {
        'color': False,
        'fon_color':False,
        'font_size': False,
        'spacing': False,
        'scale': False,
        'shading': False
    }
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if run.font.color.rgb != (0, 0, 0):
                methods['color'] = True

            if run.font.size != 152400:
                methods['font_size'] = True

            if run_get_scale(run):
                methods['scale'] = True

            if run_get_spacing(run):
                methods['spacing'] = True

            if paragraph.style.font.color.rgb != (255, 255, 255):  
                methods['shading'] = True
 
    return methods


def determine_hiding_method(methods):
    for method, value in methods.items():
        if value:
            return method
    return None


def run_get_spacing(run):
    rPr = run._r.get_or_add_rPr()
    spacings = rPr.xpath("./w:spacing")
    return spacings


def run_get_scale(run):
    rPr = run._r.get_or_add_rPr()
    spacings = rPr.xpath("./w:w")
    return spacings


def encode_text(document, chosen_method):
    encoded_text = ""
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            for char in run.text:
                if chosen_method == 'color':
                    if run.font.color.rgb == (1, 1, 1):
                        encoded_text += "1"
                    else:
                        encoded_text += "0"

                elif chosen_method == 'font_size':
                    if run.font.size != 152400:
                        encoded_text += "1"
                    else:
                        encoded_text += "0"
                        
                elif chosen_method == 'spacing':
                    text_array_t = []
                    text_array_t.append(run_get_spacing(run))
                    
                    if (text_array_t[0]) == []:
                        encoded_text += "0"
                    else:
                        encoded_text += "1"
                        
                elif chosen_method == 'scale':
                    text_array_t = []
                    text_array_t.append(run_get_scale(run))

                    if (text_array_t[0]) == []:
                        encoded_text += "0"
                    else:
                        encoded_text += "1"

                else:
                    encoded_text += "0"  # В случае, если метод не определен
    return encoded_text


def decode_text(encoded_text):
    decoded_results = []
    chunks = [chunk for chunk in encoded_text.split("000000000") if chunk and chunk != '000000000']
    for index, chunk in enumerate(chunks):
        # Пропустить последний чанк, состоящий только из нулей
        if index == len(chunks) - 1:
            break
        chunk += "0" * (-len(chunk) % 8)  # Дополним последний блок нулями до длины, кратной 8
        binary_sequence = " ".join([chunk[i:i+8] for i in range(0, len(chunk), 8)])[:-1]  # Преобразование в строку с разделением по 8 бит и удаление последних 8 символов
        decoded_results.extend(decode_text_for_encodings(binary_sequence))
    return decoded_results


def print_results(decoded_results, chosen_method):
    print(f"Метод сокрытия информации: {chosen_method if chosen_method else 'не определен'}")
    for binary_sequence, encoding, decoded_text in decoded_results:
        if decoded_text:
            print(f"Кодировка: {encoding}")
            print(f"Битовая последовательность: {binary_sequence}")
            print(f"Результат: {decoded_text[:-1]}\n")


def decode_text_for_encodings(binary_sequence):
    decoded_results = []
    encodings = ["koi8-r", "koi8-u", "cp866", "windows-1251", "cp1251", "Baudot"]  
    for encoding in encodings:
        binary_sequence, decoded_text = decode_bytes_to_string(binary_sequence, encoding)
        if decoded_text:
            decoded_results.append((binary_sequence, encoding, decoded_text))
    return decoded_results



def decode_bytes_to_string(binary_sequence, encoding):
    if encoding == "Baudot":
        return decode_baudot(binary_sequence)
    try:
        bytes_list = [int(byte, 2) for byte in binary_sequence.split()]
        bytes_data = bytes(bytes_list)
        decoded_text = bytes_data.decode(encoding)
        return binary_sequence, decoded_text
    except Exception as e:
        return None, None


def decode_baudot(binary_sequence):
    baudot_to_char = {
        "00000": " ", "00001": "А", "00010": "Б", "00011": "В", "00100": "Г",
        "00101": "Д", "00110": "Е", "00111": "Ё", "01000": "Ж", "01001": "З",
        "01010": "И", "01011": "Й", "01100": "К", "01101": "Л", "01110": "М",
        "01111": "Н", "10000": "О", "10001": "П", "10010": "Р", "10011": "С",
        "10100": "Т", "10101": "У", "10110": "Ф", "10111": "Х", "11000": "Ц",
        "11001": "Ч", "11010": "Ш", "11011": "Щ", "11100": "Ъ", "11101": "Ы",
        "11110": "Ь", "11111": "Э"
    }

    decoded_text = ""
    # Разбиваем последовательность на пятибитовые части и декодируем каждую часть
    for i in range(0, len(binary_sequence), 5):
        byte = binary_sequence[i:i+5]
        if byte in baudot_to_char:
            decoded_text += baudot_to_char[byte]
    return binary_sequence, decoded_text


# Пример использования
analyze_document("C://Users//daria//OneDrive//Рабочий стол//10 семестр//Мат медоды сокрытия и маскирования информации//лаба1//variant04.docx")
